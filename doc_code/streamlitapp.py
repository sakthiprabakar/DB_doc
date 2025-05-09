import streamlit as st
import re
import pandas as pd
import os
import json
from io import StringIO, BytesIO
import boto3
from botocore.exceptions import ClientError
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time
import botocore.config

# Set page configuration
st.set_page_config(
    page_title="SQL Stored Procedure Analyzer",
    page_icon="🧰",
    layout="wide"
)

# Function to create a Word document from analysis
def create_word_document(analysis):
    # Create a new Document
    doc = Document()

    # Add title
    title = doc.add_heading('SQL Stored Procedure Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add procedure name
    proc_name_heading = doc.add_heading('Stored Procedure Name:', level=1)
    # Make the procedure name itself bold in its own paragraph
    proc_name_para = doc.add_paragraph()
    proc_name_para.add_run(analysis['procedure_name']).bold = True

    # Add complexity
    doc.add_heading('Complexity:', level=1)
    doc.add_paragraph(analysis['complexity'])

    # Add scope
    doc.add_heading('Scope:', level=1)
    doc.add_paragraph(analysis['scope'])

    # Add optimization steps
    doc.add_heading('Optimization Steps:', level=1)

    for i, opt in enumerate(analysis["optimizations"], 1):
        # Step heading
        step_heading = doc.add_heading(f'Step {i}: {opt["type"]}', level=2)

        # Existing Logic
        doc.add_heading('Existing Logic:', level=3)
        existing_code = doc.add_paragraph(opt["existing_logic"])

        # Format code paragraph
        existing_code_fmt = existing_code.paragraph_format
        existing_code_fmt.left_indent = Inches(0.25)
        existing_code_fmt.right_indent = Inches(0.25)
        # Apply monospaced font to the runs within the paragraph
        for run in existing_code.runs:
            run.font.name = 'Courier New'
            # Optional: Set font size for code
            run.font.size = Pt(10)

        # Optimized Logic
        doc.add_heading('Optimized Logic:', level=3)
        optimized_code = doc.add_paragraph(opt["optimized_logic"])

        # Format code paragraph
        optimized_code_fmt = optimized_code.paragraph_format
        optimized_code_fmt.left_indent = Inches(0.25)
        optimized_code_fmt.right_indent = Inches(0.25)
        # Apply monospaced font to the runs within the paragraph
        for run in optimized_code.runs:
            run.font.name = 'Courier New'
            # Optional: Set font size for code
            run.font.size = Pt(10)

        # Explanation
        explanation_para = doc.add_paragraph()
        explanation_text = explanation_para.add_run(opt["explanation"])
        explanation_text.italic = True

        # Add separator paragraph
        separator = doc.add_paragraph()
        separator.add_run('_' * 40)

    # Add summary table
    doc.add_heading('Summary:', level=1)

    # Create table data (ensure line_number is handled if missing)
    table_data = []
    for opt in analysis["optimizations"]:
        table_data.append({
            "Type of Change": opt.get("type", "N/A"),
            "Line Number": opt.get("line_number", "N/A"), # Use .get for safety
            "Original Code Snippet": opt.get("existing_logic", ""),
            "Optimized Code Snippet": opt.get("optimized_logic", ""),
            "Optimization Explanation": opt.get("explanation", "")
        })

    # Add table to document only if there's data
    if table_data:
        num_rows = 1 + len(table_data)
        num_cols = 5
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid' # Ensure 'Table Grid' style exists or use a known default

        # Set header row
        header_cells = table.rows[0].cells
        headers = ['Type of Change', 'Line Number', 'Original Code Snippet', 'Optimized Code Snippet', 'Optimization Explanation']
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            # Clear existing content (sometimes needed)
            cell.text = ''
            # Add text and format
            p = cell.paragraphs[0]
            run = p.add_run(header_text)
            run.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER # Center align header text

        # Add data rows
        for i, item in enumerate(table_data):
            row_cells = table.rows[i+1].cells
            row_cells[0].text = item['Type of Change']
            row_cells[1].text = item['Line Number']
            row_cells[2].text = item['Original Code Snippet']
            row_cells[3].text = item['Optimized Code Snippet']
            row_cells[4].text = item['Optimization Explanation']

        # Set table column widths
        try:
            table.columns[0].width = Inches(1.2)
            table.columns[1].width = Inches(0.8) # Reduced width for line number
            table.columns[2].width = Inches(1.5)
            table.columns[3].width = Inches(1.5)
            table.columns[4].width = Inches(2.0) # Increased width for explanation
        except IndexError:
             st.warning("Could not set all table column widths.")

        # Apply alternate row shading
        for i, row in enumerate(table.rows):
            if i > 0 and i % 2 == 0:  # Even data rows (index 2, 4, ...) get shading
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    shading_elm = OxmlElement('w:shd')
                    shading_elm.set(qn('w:fill'), "F2F2F2") # Light gray color
                    shading_elm.set(qn('w:val'), 'clear') # Ensure fill type is set
                    shading_elm.set(qn('w:color'), 'auto')
                    tcPr.append(shading_elm)
    else:
        doc.add_paragraph("No optimization suggestions were generated.")

    # Save the document to a BytesIO object
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    return doc_io

# Function to analyze stored procedure using AWS Bedrock with Claude 3.5 Sonnet
def analyze_stored_procedure(file_content):
    try:
        # Get AWS credentials from streamlit secrets
        aws_access_key = st.secrets["aws"]["aws_access_key_id"]
        aws_secret_key = st.secrets["aws"]["aws_secret_access_key"]
        aws_region = st.secrets["aws"]["aws_region"]
        
        # Initialize AWS Bedrock client with increased timeout and retry configuration
        boto_config = botocore.config.Config(
            connect_timeout=30,  # Increase connection timeout to 30 seconds
            read_timeout=120,    # Increase read timeout to 120 seconds
            retries={
                'max_attempts': 5,  # Maximum number of retry attempts
                'mode': 'standard'  # Use standard retry mode
            }
        )
        
        bedrock = boto3.client(
            'bedrock-runtime', 
            region_name=aws_region,
            aws_access_key_id=aws_access_key,
            aws_secret_access_key=aws_secret_key,
            config=boto_config
        )
        model_id = "anthropic.claude-3-5-sonnet-20240620-v1:0"
        
        # Create prompt for analysis with improved formatting instructions
        prompt = f"""
        Analyze the following SQL stored procedure in its entirety and return your analysis in JSON format:
        
        {file_content}
        
        You are an expert in SQL performance optimization. Thoroughly analyze the stored procedure and identify the most significant optimization opportunities specific to this code. Focus on optimizations that would result in meaningful performance improvements.
        
        Extract and provide:
        1. The name of the stored procedure
        2. The complexity level of the query that needs to be optimized
        3. The scope/purpose of the stored procedure with details of 4-5 lines
        4. Key optimization opportunities

        For each optimization opportunity:
        - Provide a clear, descriptive name for the type of optimization
        - IMPORTANT: Indicate the EXACT line number range in the code where this optimization applies (e.g., "15-20", "32-45")
          Do NOT use generic terms like "Entire procedure" or "Throughout the procedure"
        - Include the existing code snippet (SIMPLIFIED AND SANITIZED for JSON)
        - Provide the improved code snippet (SIMPLIFIED AND SANITIZED for JSON)
        - Explain the performance benefit and why this change would be impactful
        
        EXTREMELY IMPORTANT JSON FORMATTING INSTRUCTIONS:
        1. ALL CODE SNIPPETS MUST BE SIMPLE TEXT WITHOUT SPECIAL FORMATTING
        2. Replace all double quotes in code with single quotes
        3. Replace all newlines in code with the literal string "\\n"
        4. Replace all tabs with the literal string "\\t"
        5. Replace all backslashes with double backslashes "\\\\"
        6. DO NOT include any raw newlines, tabs, or unescaped quotes in JSON values
        7. Keep code examples simple and focus on the key changes
        
        Output format:
        {{
            "procedure_name": "name of the procedure",
            "complexity": "complexity level",
            "scope": "brief description",
            "optimizations": [
                {{
                    "type": "type of optimization",
                    "line_number": "specific line number range (e.g., '15-20')",
                    "existing_logic": "simplified code with newlines as \\n",
                    "optimized_logic": "simplified code with newlines as \\n",
                    "explanation": "explanation of benefits"
                }}
            ],
            "summary": {{
                "original_performance_issues": "key issues overview",
                "optimization_impact": "estimated impact",
                "implementation_difficulty": "difficulty assessment"
            }}
        }}
        
        FINAL INSTRUCTIONS:
        1. Your response must contain ONLY valid JSON.
        2. Do NOT include backticks or JSON code block markers.
        3. All string values must be properly escaped for JSON.
        4. Use simple, sanitized code examples without complex formatting.
        5. Double-check that your JSON response will parse correctly before returning it.
        6. For each optimization, always provide specific line number ranges, never general locations.

        Consider these optimization strategies if applicable:
        1) Index usage optimization
        2) Remove redundant DISTINCT/UNION operations
        3) Replace row-by-row processing with set-based operations
        4) Consolidate NOCOUNT usage
        5) Use CTEs instead of nested loops for parsing
        6) Replace SELECT * with specific columns
        7) Minimize dynamic SQL usage
        8) Use temporary tables instead of cursors
        """
        
        # Generate response using Claude model with retry logic
        def generate_response(prompt):
            payload = {
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": 8000,
                "temperature": 0,
                "messages": [
                    {
                        "role": "user",
                        "content": [{"type": "text", "text": prompt}],
                    }
                ],
                "system": "You are an expert SQL database optimizer. Your responses must be valid, properly escaped JSON without any special characters or line breaks in JSON string values. Always format code snippets by replacing newlines with \\n, tabs with \\t, and using single quotes instead of double quotes whenever possible."
            }
            
            # Initialize retry variables
            max_retries = 3
            retry_count = 0
            retry_delay = 5  # Initial delay of 5 seconds
            
            while retry_count <= max_retries:
                try:
                    response = bedrock.invoke_model(
                        modelId=model_id,
                        body=json.dumps(payload),
                        accept="application/json",
                        contentType="application/json"
                    )
                    
                    response_body = json.loads(response["body"].read())
                    return response_body["content"][0]["text"]
                    
                except (ClientError, Exception) as e:
                    retry_count += 1
                    error_message = str(e)
                    
                    # If we've exhausted all retries, give up and report the error
                    if retry_count > max_retries:
                        st.error(f"ERROR: Can't invoke model after {max_retries} attempts. Reason: {error_message}")
                        return None
                    
                    # Log the retry attempt
                    st.warning(f"Retry {retry_count}/{max_retries}: Error invoking model. Reason: {error_message}")
                    st.info(f"Waiting {retry_delay} seconds before retrying...")
                    
                    # Wait before retrying with exponential backoff
                    time.sleep(retry_delay)
                    retry_delay *= 2  # Exponential backoff
        
        # Get analysis results
        analysis_result = generate_response(prompt)
        
        if not analysis_result:
            return None
            
        # Debug: Display raw response for troubleshooting
        st.sidebar.expander("Debug Raw Response", expanded=False).code(analysis_result)
        
        # Clean the response: Remove any markdown formatting if present
        cleaned_response = analysis_result.strip()
        if cleaned_response.startswith("```json"):
            cleaned_response = cleaned_response[7:]  # Remove ```json prefix
        if cleaned_response.endswith("```"):
            cleaned_response = cleaned_response[:-3]  # Remove ``` suffix
        
        # Add additional pre-processing to fix common JSON issues
        cleaned_response = fix_json_formatting(cleaned_response)
            
        # Parse the JSON with additional error handling and repair attempts
        try:
            # First attempt: Try direct JSON parse
            analysis_json = json.loads(cleaned_response)
            
            # Validate the structure (basic check to ensure we have the expected fields)
            if not all(key in analysis_json for key in ["procedure_name", "complexity", "scope", "optimizations", "summary"]):
                st.warning("JSON parsed successfully but missing required fields. Attempting to repair...")
                raise json.JSONDecodeError("Missing required fields", cleaned_response, 0)
                
            return analysis_json
            
        except json.JSONDecodeError as e:
            st.warning(f"Initial JSON parsing failed: {str(e)}. Attempting to repair...")
            
            # Try the advanced repair function
            repaired_json = repair_json_response(cleaned_response, str(e))
            if repaired_json:
                return repaired_json
                
            # If repair failed, return a minimal functional structure
            return create_fallback_response(cleaned_response)
    
    except Exception as e:
        st.error(f"Error during analysis: {str(e)}")
        import traceback
        st.sidebar.expander("Error Details", expanded=False).code(traceback.format_exc())
        return None


def fix_json_formatting(json_text):
    """Pre-process JSON to fix common formatting issues before parsing."""
    # Replace any literal \n that should be escaped newlines
    json_text = re.sub(r'([^\\])\\n', r'\1\\n', json_text)
    
    # Replace any literal \t that should be escaped tabs
    json_text = re.sub(r'([^\\])\\t', r'\1\\t', json_text)
    
    # Fix any unescaped quotes in string values
    lines = json_text.split('\n')
    in_string = False
    fixed_lines = []
    
    for line in lines:
        fixed_line = ""
        i = 0
        while i < len(line):
            if line[i] == '"' and (i == 0 or line[i-1] != '\\'):
                in_string = not in_string
                fixed_line += '"'
            elif line[i] == '"' and line[i-1] == '\\' and line[i-2] == '\\':
                # This is an escaped backslash followed by a quote, not an escaped quote
                in_string = not in_string
                fixed_line += '"'
            elif in_string and line[i] == '"' and line[i-1] != '\\':
                # Unescaped quote inside a string - escape it
                fixed_line += '\\"'
            else:
                fixed_line += line[i]
            i += 1
        
        fixed_lines.append(fixed_line)
    
    return '\n'.join(fixed_lines)


def repair_json_response(json_text, error_message):
    """Advanced JSON repair function."""
    try:
        # Extract useful information from error message
        if "Unterminated string" in error_message:
            # Find the position of the error
            match = re.search(r'line (\d+) column (\d+)', error_message)
            if match:
                line_num = int(match.group(1))
                column_num = int(match.group(2))
                
                # Split the JSON text into lines
                lines = json_text.split('\n')
                
                # If the error is within range
                if 0 <= line_num - 1 < len(lines):
                    problematic_line = lines[line_num - 1]
                    
                    # Fix unterminated string by adding a closing quote
                    # This is a simplistic approach - a more sophisticated approach would check string balance
                    if column_num - 1 < len(problematic_line):
                        # Count quotes before the error position to determine if we need to add a closing quote
                        quote_count = problematic_line[:column_num].count('"') - problematic_line[:column_num].count('\\"')
                        
                        if quote_count % 2 == 1:  # Odd number of quotes means we need a closing quote
                            lines[line_num - 1] = problematic_line[:column_num] + '"' + problematic_line[column_num:]
                            
                            # Rebuild the JSON text
                            fixed_json = '\n'.join(lines)
                            try:
                                return json.loads(fixed_json)
                            except json.JSONDecodeError:
                                st.warning("Failed to fix unterminated string")
        
        # Try using a more lenient JSON parser if available
        try:
            import demjson3
            return demjson3.decode(json_text)
        except (ImportError, Exception):
            st.warning("demjson3 parser not available or failed")
        
        # Extract complete JSON object using regex
        json_pattern = re.compile(r'\{.*\}', re.DOTALL)
        match = json_pattern.search(json_text)
        if match:
            extracted_json = match.group(0)
            try:
                return json.loads(extracted_json)
            except json.JSONDecodeError:
                st.warning("Regex extraction failed")
        
        # If all else fails, try the json5 library if available
        try:
            import json5
            return json5.loads(json_text)
        except (ImportError, Exception):
            st.warning("json5 parser not available or failed")
            
        return None
        
    except Exception as e:
        st.warning(f"JSON repair attempt failed: {str(e)}")
        return None


def create_fallback_response(json_text):
    """Create a minimal functional JSON response when parsing fails."""
    # Extract procedure name if possible
    proc_name_match = re.search(r'"procedure_name"\s*:\s*"([^"]+)"', json_text)
    proc_name = proc_name_match.group(1) if proc_name_match else "Unknown"
    
    # Extract complexity if possible
    complexity_match = re.search(r'"complexity"\s*:\s*"([^"]+)"', json_text)
    complexity = complexity_match.group(1) if complexity_match else "Medium"
    
    # Extract scope if possible
    scope_match = re.search(r'"scope"\s*:\s*"([^"]+)"', json_text)
    scope = scope_match.group(1) if scope_match else "This stored procedure's scope could not be determined due to parsing issues."
    
    # Create a minimal functional JSON response
    minimal_json = {
        "procedure_name": proc_name,
        "complexity": complexity,
        "scope": scope,
        "optimizations": [
            {
                "type": "General Optimization",
                "line_number": "N/A",
                "existing_logic": "-- Original procedure code (could not be parsed)",
                "optimized_logic": "-- See recommendations in the AI analysis text",
                "explanation": "The JSON response from the AI service could not be fully parsed. Please review the raw response in the debug section."
            }
        ],
        "summary": {
            "original_performance_issues": "The JSON response could not be fully parsed.",
            "optimization_impact": "See debug output for details.",
            "implementation_difficulty": "N/A"
        }
    }
    
    st.warning("Created fallback response due to parsing issues. Check debug output for raw response.")
    return minimal_json

# UI Components 
st.title("SQL Stored Procedure Analyzer")
st.write("Upload a SQL stored procedure file for AI-powered optimization analysis")

st.markdown("---")
st.markdown("""
    ### About This Tool
    This app analyzes SQL stored procedures using AI to identify optimization opportunities.
    
    **Features:**
    - Extract procedure name and purpose
    - Identify optimization opportunities
    - Generate improved SQL code
    - Provide a summary of changes
    - Download formatted report as Word document
    """)

# Sample SQL button for testing
if st.button("Load Sample SQL for Testing"):
    sample_sql = """
    CREATE PROCEDURE usp_GetCustomerOrders
    @CustomerId INT
    AS
    BEGIN
        SET NOCOUNT ON;
        
        -- Create temp table to store order data
        CREATE TABLE #TempOrders (
            OrderId INT,
            OrderDate DATETIME,
            OrderAmount DECIMAL(18,2)
        )
        
        -- Insert data into temp table
        INSERT INTO #TempOrders
        SELECT 
            OrderId,
            OrderDate,
            OrderAmount
        FROM Orders
        WHERE CustomerId = @CustomerId
        
        -- Cursor to process orders
        DECLARE @OrderId INT
        DECLARE @OrderDate DATETIME
        
        DECLARE order_cursor CURSOR FOR
        SELECT OrderId, OrderDate FROM #TempOrders
        
        OPEN order_cursor
        FETCH NEXT FROM order_cursor INTO @OrderId, @OrderDate
        
        WHILE @@FETCH_STATUS = 0
        BEGIN
            -- Update order status
            UPDATE Orders SET Status = 'Processed' WHERE OrderId = @OrderId
            UPDATE Orders SET LastModified = GETDATE() WHERE OrderId = @OrderId
            
            -- Process order details
            UPDATE OrderDetails 
            SET Processed = 1 
            WHERE OrderId = @OrderId
            
            FETCH NEXT FROM order_cursor INTO @OrderId, @OrderDate
        END
        
        CLOSE order_cursor
        DEALLOCATE order_cursor
        
        -- Return results
        SELECT 
            c.CustomerName,
            o.OrderId,
            o.OrderDate,
            o.OrderAmount,
            (SELECT COUNT(*) FROM OrderDetails WHERE OrderId = o.OrderId) AS ItemCount
        FROM 
            Customers c
            INNER JOIN Orders o ON c.CustomerId = o.CustomerId
        WHERE 
            c.CustomerId = @CustomerId
            
        -- Cleanup
        DROP TABLE #TempOrders
    END
    """
    st.session_state['sample_sql'] = sample_sql
    st.success("Sample SQL loaded! Click 'Analyze' to process it.")

# File upload component
uploaded_file = st.file_uploader("Upload SQL Stored Procedure", type=["sql"])

# Get SQL either from upload or sample
sql_content = None
if uploaded_file:
    sql_content = uploaded_file.getvalue().decode("utf-8")
elif 'sample_sql' in st.session_state:
    sql_content = st.session_state['sample_sql']
    st.info("Using sample SQL procedure. You can upload your own file to replace it.")

if sql_content:
    # Display the SQL
    with st.expander("View SQL Code", expanded=False):
        st.code(sql_content, language="sql")
    
    # Analysis button
    if st.button("Analyze SQL Procedure"):
        # Run analysis
        with st.spinner("Analyzing stored procedure... This may take up to 60 seconds."):
            analysis = analyze_stored_procedure(sql_content)
        
        if analysis:
            # Display results in tabs
            tab1, tab2 = st.tabs(["Analysis", "Download Report"])
            
            with tab1:
                # Display the procedure name and scope
                st.header(f"🔹 Stored Proc Name: `{analysis['procedure_name']}`")
                
                # Display the complexity
                st.subheader("🔹 Complexity:")
                st.write(analysis["complexity"])
                
                st.subheader("🔹 Scope:")
                st.write(analysis["scope"])
                
                # Display optimization steps
                st.subheader("🔹 Optimization Steps:")
                
                for i, opt in enumerate(analysis["optimizations"], 1):
                    st.markdown(f"⚙️ **Step {i}**: {opt['type']}")
                    
                    st.markdown("**Existing Logic:**")
                    st.code(opt["existing_logic"], language="sql")
                    
                    st.markdown("**Optimized Logic:**")
                    st.code(opt["optimized_logic"], language="sql")
                    
                    st.markdown(f"*{opt['explanation']}*")
                    st.markdown("---")
                
                # Create and display summary table
                st.subheader("🔹 Summary:")
                
                summary_data = []
                for opt in analysis["optimizations"]:
                    summary_data.append({
                        "Type of Change": opt["type"],
                        "Line Number": opt["line_number"],
                        "Original Code Snippet": opt["existing_logic"],
                        "Optimized Code Snippet": opt["optimized_logic"],
                        "Optimization Explanation": opt["explanation"]
                    })
                
                summary_df = pd.DataFrame(summary_data)
                
                # Display as a formatted table with custom styling
                st.markdown("""
                <style>
                .summary-table {
                    font-size: 0.85rem;
                    border-collapse: collapse;
                    width: 100%;
                }
                .summary-table th {
                    background-color: #f2f2f2;
                    text-align: left;
                    padding: 8px;
                    border: 1px solid #ddd;
                }
                .summary-table td {
                    text-align: left;
                    padding: 8px;
                    border: 1px solid #ddd;
                }
                .summary-table tr:nth-child(even) {
                    background-color: #f9f9f9;
                }
                </style>
                """, unsafe_allow_html=True)
                
                # Convert dataframe to HTML table with custom classes
                table_html = summary_df.to_html(classes='summary-table', escape=False, index=False)
                st.markdown(table_html, unsafe_allow_html=True)
                
                # Display additional summary information
                if "summary" in analysis:
                    st.subheader("🔹 Overall Assessment:")
                    if "original_performance_issues" in analysis["summary"]:
                        st.markdown(f"**Original Issues:** {analysis['summary']['original_performance_issues']}")
                    if "optimization_impact" in analysis["summary"]:
                        st.markdown(f"**Impact of Optimizations:** {analysis['summary']['optimization_impact']}")
                    if "implementation_difficulty" in analysis["summary"]:
                        st.markdown(f"**Implementation Difficulty:** {analysis['summary']['implementation_difficulty']}")
            
            with tab2:
                # Create Word document
                with st.spinner("Generating Word document..."):
                    docx_bytes = create_word_document(analysis)
                
                # Provide download button for DOCX
                st.download_button(
                    label="⬇️ Download Report as Word Document",
                    data=docx_bytes,
                    file_name=f"{analysis['procedure_name']}_analysis.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key='docx-download'
                )
                
                # Also provide markdown option
                report_md = f"""# SQL Stored Procedure Analysis Report

## Procedure Name: {analysis['procedure_name']}

## Complexity:
{analysis['complexity']}

## Scope:
{analysis['scope']}

## Optimization Steps:
"""
                
                for i, opt in enumerate(analysis["optimizations"], 1):
                    report_md += f"""
### Step {i}: {opt['type']}

**Existing Logic:**
```sql
{opt['existing_logic']}
```

**Optimized Logic:**
```sql
{opt['optimized_logic']}
```

*{opt['explanation']}*

---
"""
                
                report_md += "\n## Summary Table:\n\n"
                report_md += summary_df.to_markdown(index=False)
                
                st.download_button(
                    label="⬇️ Download Report as Markdown",
                    data=report_md,
                    file_name=f"{analysis['procedure_name']}_analysis.md",
                    mime="text/markdown",
                    key='md-download'
                )
                
                st.info("The Word document (.docx) contains the same content as shown in the 'Analysis' tab, but in a properly formatted document for sharing.")
        else:
            st.error("Analysis could not be completed. Please check the Debug section in the sidebar for more details.")

else:
    # Show example when no file is uploaded
    st.info("Please upload a SQL stored procedure file (.sql) or use the sample SQL to begin analysis.")
    
    with st.expander("See Example Analysis"):
        st.markdown("""
        ## Example Output
        
        🔹 **Stored Proc Name:** `usp_get_customer_data`
        
        🔹 **Complexity:** Medium
        
        🔹 **Scope:**  
        This procedure retrieves customer data including their purchase history, last login time, and calculates their loyalty score using internal metrics.
        
        🔹 **Optimization Steps:**
        
        ⚙️ **Step 1:** Replace Multiple Updates
        
        **Existing Logic:**
        ```sql
        UPDATE table SET col1 = val WHERE condition;
        UPDATE table SET col2 = val WHERE condition;
        ```
        
        **Optimized Logic:**
        ```sql
        UPDATE table 
        SET col1 = val, 
            col2 = val 
        WHERE condition;
        ```
        
        *Reduces write operations and improves efficiency.*
        """)
        
        # Example of the summary table
        st.markdown("🔹 **Summary:**")
        
        example_data = [{
            "Type of Change": "Replace Multiple Updates",
            "Line Number": "Identified in multiple places",
            "Original Code Snippet": "UPDATE table SET col1 = val WHERE condition;\nUPDATE table SET col2 = val WHERE condition;",
            "Optimized Code Snippet": "UPDATE table \nSET col1 = val, \n    col2 = val \nWHERE condition;",
            "Optimization Explanation": "Reduces write operations and improves efficiency."
        }, {
            "Type of Change": "Index on Temp Tables",
            "Line Number": "Where temp tables are created",
            "Original Code Snippet": "CREATE TABLE #temp (\n    id INT,\n    value VARCHAR(50)\n)",
            "Optimized Code Snippet": "CREATE TABLE #temp (\n    id INT,\n    value VARCHAR(50)\n);\nCREATE INDEX IX_temp_id ON #temp(id);",
            "Optimization Explanation": "Improves performance by speeding up lookups and joins."
        }]
        
        example_df = pd.DataFrame(example_data)
        
        # Display example table with styling
        st.markdown("""
        <style>
        .summary-table {
            font-size: 0.85rem;
            border-collapse: collapse;
            width: 100%;
        }
        .summary-table th {
            background-color: #f2f2f2;
            text-align: left;
            padding: 8px;
            border: 1px solid #ddd;
        }
        .summary-table td {
            text-align: left;
            padding: 8px;
            border: 1px solid #ddd;
        }
        .summary-table tr:nth-child(even) {
            background-color: #f9f9f9;
        }
        </style>
        """, unsafe_allow_html=True)
        
        table_html = example_df.to_html(classes='summary-table', escape=False, index=False)
        st.markdown(table_html, unsafe_allow_html=True)

# Add footer
st.markdown("---")
st.caption("SQL Stored Procedure Analyzer powered by AWS Bedrock Claude 3.5 Sonnet")